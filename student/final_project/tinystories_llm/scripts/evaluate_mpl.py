import argparse
import json
import random
import re
import os
from dataclasses import dataclass
from typing import List, Dict, Tuple
import torch

from bpe_tokenizer import BPETokenizer
from transformer_model import TinyStoriesConfig, TinyStoriesForCausalLM
from docx import Document


# Options A and B are randomized to prevent position bias

def build_model(tokenizer, max_seq_len: int) -> TinyStoriesForCausalLM:
    config = TinyStoriesConfig(
        vocab_size=len(tokenizer.token2id),
        hidden_size=256,
        num_hidden_layers=4,
        num_attention_heads=8,
        intermediate_size=1024,
        hidden_dropout_prob=0.1,
        attention_probs_dropout_prob=0.1,
        max_position_embeddings=max_seq_len,
        window_size=max_seq_len,
    )
    return TinyStoriesForCausalLM(config)


def load_model_and_tokenizer(model_path: str, tokenizer_path: str, device, max_seq_len: int):
    tokenizer = BPETokenizer.load(tokenizer_path)
    model = build_model(tokenizer, max_seq_len=max_seq_len).to(device)
    state = torch.load(model_path, map_location=device, weights_only=True)
    model.load_state_dict(state)
    model.eval()
    return model, tokenizer


def greedy_generate_until_choice_reason(model, tok, prompt, device, max_new_tokens=40, max_seq_len=256):
    text = f"<user> {prompt} <assistant>"
    prompt_ids = tok.encode(text, add_special_tokens=False)
    x = torch.tensor([prompt_ids], dtype=torch.long, device=device)

    generated = []
    eos_id = tok.token2id.get("<eos>", None)
    user_id = tok.token2id.get("<user>", None)

    with torch.no_grad():
        for _ in range(max_new_tokens):
            attn = torch.ones_like(x)
            out = model(input_ids=x, attention_mask=attn)
            nxt = int(out["logits"][0, -1].argmax().item())

            generated.append(nxt)
            x = torch.cat([x, torch.tensor([[nxt]], device=device)], dim=1)
            gen_text = tok.decode(generated)

            if user_id is not None and nxt == user_id:
                break
            if eos_id is not None and nxt == eos_id:
                break
            if x.size(1) >= max_seq_len:
                break

            if re.search(r"Choice:\s*[AB].*Reason:.*[\.\n]", gen_text, re.IGNORECASE | re.DOTALL):
                break

    return tok.decode(generated).strip()


def extract_choice(text: str) -> str:
    m = re.search(r"Choice:\s*([AB])", text, re.IGNORECASE)
    if m:
        return m.group(1).upper()
    return "N/A"


@dataclass
class SimplifiedMPLBlock:
    high_coins: int
    low_coins: int
    sure_coins: int
    probs: List[float]
    name: str

def make_fairy_tale_blocks() -> List[SimplifiedMPLBlock]:
    probs = [i / 10 for i in range(1, 11)]  # 10% to 100%

    return [
        SimplifiedMPLBlock(
            high_coins=30,
            low_coins=0,
            sure_coins=10,
            probs=probs,
            name="Small_Stakes_30_vs_10"
        ),
        SimplifiedMPLBlock(
            high_coins=100,
            low_coins=0,
            sure_coins=10,
            probs=probs,
            name="Medium_Stakes_100_vs_10"
        ),
        SimplifiedMPLBlock(
            high_coins=300,
            low_coins=0,
            sure_coins=10,
            probs=probs,
            name="High_Stakes_300_vs_10"
        ),
    ]

def format_mpl_prompt(row_idx: int, total_rows: int, p: float, block: SimplifiedMPLBlock, rng: random.Random) -> Tuple[str, str, str]:
    """
    Randomly assigns Gamble to A or B, just like the training data.
    Returns: (Prompt, Gamble_Option_Letter, Safe_Option_Letter)
    """
    p_high = int(round(p * 100))
    p_low = 100 - p_high

    gamble_text = f"{p_high}% chance to get {block.high_coins} shiny coins, and {p_low}% chance to get {block.low_coins} shiny coins."
    safe_text = f"get {block.sure_coins} shiny coins for sure."

    if rng.choice([True, False]):
        opt_a_text, opt_b_text = gamble_text, safe_text
        gamble_opt, safe_opt = "A", "B"
    else:
        opt_a_text, opt_b_text = safe_text, gamble_text
        gamble_opt, safe_opt = "B", "A"

    prompt = (
        f"Once upon a time, Tom was in a small forest. Tom wanted to trade shiny coins for a toy. "
        f"Then Tom met a merchant with a big hat. Tom felt a little nervous. "
        f"Option A: {opt_a_text} Option B: {opt_b_text} "
        f"What should Tom choose? Answer with exactly:\n"
        f"Choice: A or Choice: B\n"
        f"Reason: one short sentence."
    )
    return prompt, gamble_opt, safe_opt


def evaluate_one_block(model, tokenizer, device, block, max_new_tokens, max_seq_len, rng, debug=False):
    choices_made = []
    actions = []
    raws = []

    for idx, p in enumerate(block.probs, start=1):#for probability 0.1 to 1
        prompt, gamble_opt, safe_opt = format_mpl_prompt(idx, len(block.probs), p, block, rng)

        
        raw = greedy_generate_until_choice_reason(
            model=model, tok=tokenizer, prompt=prompt, device=device,
            max_new_tokens=max_new_tokens, max_seq_len=max_seq_len
        ) 
        
        choice = extract_choice(raw)
        choices_made.append(choice)
        raws.append(raw)

       
        if choice == safe_opt:
            action = "Safe"
        elif choice == gamble_opt:
            action = "Gamble"
        else:
            action = "Invalid"
        actions.append(action)

        if debug:
            print(f"\n--- ROW {idx} | p={p} ---")
            print(f"Gamble is {gamble_opt}, Safe is {safe_opt}")
            print(f"Model chose: {choice} -> {action}")
            print(raw)

    n_gamble = sum(1 for a in actions if a == "Gamble")
    n_safe = sum(1 for a in actions if a == "Safe")

    return {
        "block_name": block.name,
        "probs": [int(round(x * 100)) for x in block.probs],
        "choices_made": choices_made,
        "actions": actions,
        "num_Gamble": n_gamble,
        "num_Safe": n_safe,
        "raw_outputs": raws if debug else None,
    }

def save_all_blocks_to_word(results, filename="mpl_results_all_blocks.docx"):
    """
    results: List of tuples [(res_c, res_b), (res_c, res_b), ...] for each block
    """
    doc = Document()
    doc.add_heading("MPL Results (All Blocks)", level=1)

    for (res_c, res_b) in results:
        doc.add_heading(f"Block: {res_c['block_name']}", level=2)

        doc.add_paragraph(f"Cautious: Safe={res_c['num_Safe']} | Gamble={res_c['num_Gamble']}")
        doc.add_paragraph(f"Bold:     Safe={res_b['num_Safe']} | Gamble={res_b['num_Gamble']}")

        table = doc.add_table(rows=11, cols=3)
        try:
            table.style = "Table Grid"
        except Exception:
            pass  

        table.rows[0].cells[0].text = "p(high)%"
        table.rows[0].cells[1].text = "Cautious (meaning)"
        table.rows[0].cells[2].text = "Bold (meaning)"

        for i in range(10):
            table.rows[i + 1].cells[0].text = str(res_c["probs"][i])
            table.rows[i + 1].cells[1].text = res_c["actions"][i]
            table.rows[i + 1].cells[2].text = res_b["actions"][i]

        doc.add_paragraph("")

    doc.save(filename)
    print(f"[Saved] {filename}")


def main():
    p = argparse.ArgumentParser()
    p.add_argument("--tokenizer_path", type=str, default="bpe_tokenizer_tinystories.pkl")
    p.add_argument("--cautious_path", type=str, default="cautious_chat_model/final_model.pth")
    p.add_argument("--bold_path", type=str, default="bold_gambler_chat_model/final_model.pth")
    p.add_argument("--max_seq_len", type=int, default=256)
    p.add_argument("--max_new_tokens", type=int, default=80)
    

    p.add_argument("--seed", type=int, default=359)

    p.add_argument("--debug", action="store_true")
    args = p.parse_args()

    print("Working directory:", os.getcwd())

    device = torch.device("cuda" if torch.cuda.is_available() else "cpu")

    print("Loading Models...")
    cautious_model, tokenizer = load_model_and_tokenizer(
        args.cautious_path, args.tokenizer_path, device, args.max_seq_len
    )
    bold_model, _ = load_model_and_tokenizer(
        args.bold_path, args.tokenizer_path, device, args.max_seq_len
    )

    rng_cautious = random.Random(args.seed)
    rng_bold = random.Random(args.seed)

    blocks = make_fairy_tale_blocks()

    print("\nStarting Randomized MPL Stress Test...")
    all_results = []

    for bi, block in enumerate(blocks, start=1):
        print("\n" + "=" * 80)
        print(f"BLOCK {bi}/{len(blocks)} | {block.name}")
        print(f"Gamble: ({block.high_coins}/{block.low_coins} coins) | Safe: ({block.sure_coins} coins for sure)")
        print("-" * 80)

        res_c = evaluate_one_block(
            cautious_model, tokenizer, device, block,
            args.max_new_tokens, args.max_seq_len, rng_cautious, args.debug
        )
        res_b = evaluate_one_block(
            bold_model, tokenizer, device, block,
            args.max_new_tokens, args.max_seq_len, rng_bold, args.debug
        )

        all_results.append((res_c, res_b))

        print(f"{'Row':<5} {'p(high)%':<10} {'Cautious Model':<25} {'Bold Model':<25}")
        print(f"{'':<5} {'':<10} {'Choice -> Meaning':<25} {'Choice -> Meaning':<25}")
        for i, pprob in enumerate(res_c["probs"], start=1):
            c_str = f"{res_c['choices_made'][i-1]} -> {res_c['actions'][i-1]}"
            b_str = f"{res_b['choices_made'][i-1]} -> {res_b['actions'][i-1]}"
            print(f"{i:<5} {pprob:<10} {c_str:<25} {b_str:<25}")

        print("-" * 80)
        print(f"[Cautious] Safe: {res_c['num_Safe']} | Gambles: {res_c['num_Gamble']}")
        print(f"[Bold]     Safe: {res_b['num_Safe']} | Gambles: {res_b['num_Gamble']}")

    save_all_blocks_to_word(all_results, filename="mpl_results_all_blocks.docx")



if __name__ == "__main__":
    main()